"""DOCX reading and writing.

The whole pipeline depends on one idea: every piece of text in the document is
yielded as a TextUnit, no matter whether it came from a body paragraph, a table
cell, or a header. Detection and replacement then treat all of them identically,
so a PII rule can never work "in paragraphs but not in tables" - a class of bug
that is easy to ship and hard to notice in a 400-page document.
"""

import base64
from dataclasses import dataclass
from typing import Iterator

from docx import Document
from docx.document import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph


@dataclass
class TextUnit:
    """One paragraph of text plus a human-readable note about where it lives.

    `paragraph` is the live python-docx object, so writing back to it edits the
    document in place and keeps the original styling.
    """

    text: str
    location: str
    paragraph: Paragraph


def _iter_paragraphs(container, location: str) -> Iterator[TextUnit]:
    """Yield paragraphs from a container, recursing into tables.

    Tables can nest inside table cells, so this recurses rather than looping
    twice over document.paragraphs and document.tables.
    """
    for para in container.paragraphs:
        if para.text.strip():
            yield TextUnit(text=para.text, location=location, paragraph=para)

    for t_idx, table in enumerate(container.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cell_loc = f"{location}/table{t_idx}[r{r_idx},c{c_idx}]"
                yield from _iter_paragraphs(cell, cell_loc)


def iter_text_units(doc: DocxDocument) -> Iterator[TextUnit]:
    """Yield every text-bearing paragraph in the document.

    Headers and footers are included because contact details and company names
    routinely live there, and they are invisible to document.paragraphs.
    """
    yield from _iter_paragraphs(doc, "body")

    for s_idx, section in enumerate(doc.sections):
        for part_name in ("header", "first_page_header", "even_page_header",
                          "footer", "first_page_footer", "even_page_footer"):
            part = getattr(section, part_name, None)
            if part is not None:
                yield from _iter_paragraphs(part, f"section{s_idx}/{part_name}")


def load_document(path: str) -> DocxDocument:
    return Document(path)


def document_from_text(text: str, path: str) -> str:
    """Wrap plain text in a DOCX so the one pipeline handles both inputs.

    Users often have a snippet of text rather than a Word file. Rather than
    building a second redaction path for strings - which would inevitably drift
    from the DOCX one - the text is converted into a real document first. There
    is then exactly one implementation to test and maintain, and the output is
    a .docx either way.

    Blank lines are preserved so pasted text keeps its shape.
    """
    doc = Document()
    for line in text.splitlines() or [""]:
        doc.add_paragraph(line)
    doc.save(path)
    return path


def save_document(doc: DocxDocument, path: str) -> None:
    doc.save(path)


def replace_paragraph_text(paragraph: Paragraph, new_text: str) -> None:
    """Overwrite a paragraph's text while keeping its formatting.

    DOCX splits text across "runs" at every formatting or spellcheck boundary,
    so a single email address is often three or four runs. Replacing text run by
    run would therefore corrupt entities that straddle a boundary.

    Instead the full replacement text is written into the first run and the
    remaining runs are emptied. The paragraph keeps the first run's formatting,
    which is the dominant style in practice. This trades away mid-paragraph
    formatting changes for correctness of the redaction, which is the right way
    round for this tool.
    """
    if not paragraph.runs:
        paragraph.text = new_text
        return

    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def count_images(doc: DocxDocument) -> int:
    """Count embedded images.

    Images matter because text extraction cannot see inside them. The reference
    prospectus hides a photographed PAN card - name, father's name, date of
    birth, PAN number - in an image, which is invisible to every text-based
    rule in this tool. Counting them lets the tool report the blind spot
    instead of silently ignoring it.
    """
    return sum(
        1 for rel in doc.part.rels.values() if "image" in rel.reltype
    )


def blank_images(doc: DocxDocument) -> int:
    """Overwrite every embedded image with a 1x1 transparent PNG.

    This is deliberately crude. Without OCR the tool cannot know which images
    contain PII, so it treats all of them as unreviewable and neutralises them.
    Replacing the image bytes keeps the document structure and layout intact,
    which editing the drawing XML would not.
    """
    blanked = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            rel.target_part._blob = _BLANK_PNG
            blanked += 1
    return blanked


# Smallest valid PNG: 1x1, fully transparent.
_BLANK_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNk"
    "YPhfDwAChwGA60e6kgAAAABJRU5ErkJggg=="
)
