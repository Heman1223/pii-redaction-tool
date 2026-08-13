"""Build the synthetic evaluation corpus and its ground truth.

Why this file exists: the real prospectus contains no SSNs, credit cards or IP
addresses, and its only date of birth is locked inside an image. Those four PII
types are required by the assignment, so without a second corpus they could
never be scored - the report would show blanks and look like broken detection.

This document deliberately mixes real PII with near-miss distractors (invoice
numbers shaped like SSNs, share counts shaped like card numbers, version
strings shaped like IPs, ordinary dates next to a real DOB) so that precision is
actually tested rather than assumed.
"""

import json
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]

# (paragraph text, [(entity_text, entity_type), ...])
# The empty list means "this line contains no PII" - those lines are the
# distractors that make the precision number meaningful.
CONTENT = [
    ("Employee Onboarding Record - Internal", []),
    ("Full Name: Rashi Patil", [("Rashi Patil", "PERSON")]),
    ("Email: rashhi.patil@gmail.com", [("rashhi.patil@gmail.com", "EMAIL")]),
    ("Contact: +91 9876543210", [("+91 9876543210", "PHONE")]),
    ("Date of Birth: 15 March 1992", [("15 March 1992", "DOB")]),
    ("SSN: 123-45-6789", [("123-45-6789", "SSN")]),
    ("Card on file: 4111 1111 1111 1111", [("4111 1111 1111 1111", "CREDIT_CARD")]),
    ("Last login from IP 192.168.14.22", [("192.168.14.22", "IP_ADDRESS")]),
    ("Employer: Acme Manufacturing Limited", [("Acme Manufacturing Limited", "COMPANY")]),
    (
        "Residence: 42 Baker Street, Kothrud, Pune - 411 038, Maharashtra, India",
        [("42 Baker Street, Kothrud, Pune - 411 038, Maharashtra, India", "ADDRESS")],
    ),

    # Second person, so consistency across repeats can be checked.
    ("Reporting Manager: Rohan Dey", [("Rohan Dey", "PERSON")]),
    ("Manager email: rohan.dey@gmail.com", [("rohan.dey@gmail.com", "EMAIL")]),
    ("Rashi Patil was confirmed by Rohan Dey on the same date.",
     [("Rashi Patil", "PERSON"), ("Rohan Dey", "PERSON")]),

    # --- Same-name edge case: two different people, distinguished by role ---
    ("Rahul Sharma, Director, Acme Manufacturing Limited, approved the request.",
     [("Rahul Sharma", "PERSON"), ("Acme Manufacturing Limited", "COMPANY")]),
    ("Rahul Sharma, Chief Financial Officer, Globex Industries Limited, countersigned.",
     [("Rahul Sharma", "PERSON"), ("Globex Industries Limited", "COMPANY")]),

    # --- Distractors: must NOT be redacted ---
    ("Invoice number 456-78-9012 was raised for this onboarding.", []),
    ("The policy was approved by the board on December 10, 2025.", []),
    ("Total shares allotted: 26704570 at face value 5 each.", []),
    ("Software build version 10.2.14.3 was deployed to production.", []),
    ("Corporate Identity Number: U28129PN1979PLC141032", []),
    ("SEBI Registration Number: INM000013004", []),
    ("The company was incorporated on July 30, 1979 under the Companies Act, 1956.", []),
    ("Reference ticket ID 998877665544332211 remains open.", []),
]

# Rows go into a table, which proves the pipeline reads table cells and not just
# document.paragraphs.
TABLE_ROWS = [
    ("Name", "Role", "Email", "Phone"),
    ("Priya Nair", "Company Secretary", "priya.nair@example.org", "+91 20 45053237"),
    ("Arjun Mehta", "Managing Director", "arjun.mehta@example.org", "+91 22 4009 4400"),
]

TABLE_GROUND_TRUTH = [
    ("Priya Nair", "PERSON"), ("priya.nair@example.org", "EMAIL"),
    ("+91 20 45053237", "PHONE"),
    ("Arjun Mehta", "PERSON"), ("arjun.mehta@example.org", "EMAIL"),
    ("+91 22 4009 4400", "PHONE"),
]


def build() -> None:
    doc = Document()
    doc.add_heading("Synthetic PII Test Document", level=1)

    ground_truth = []

    for text, entities in CONTENT:
        doc.add_paragraph(text)
        for value, etype in entities:
            ground_truth.append({"text": value, "type": etype})

    doc.add_heading("Contact Directory (table)", level=2)
    table = doc.add_table(rows=len(TABLE_ROWS), cols=4)
    table.style = "Table Grid"
    for r, row in enumerate(TABLE_ROWS):
        for c, value in enumerate(row):
            table.cell(r, c).text = value

    for value, etype in TABLE_GROUND_TRUTH:
        ground_truth.append({"text": value, "type": etype})

    (ROOT / "input").mkdir(exist_ok=True)
    (ROOT / "data").mkdir(exist_ok=True)

    doc_path = ROOT / "input" / "synthetic_sample.docx"
    gt_path = ROOT / "data" / "ground_truth_synthetic.json"

    doc.save(doc_path)
    gt_path.write_text(json.dumps(ground_truth, indent=2), encoding="utf-8")

    print(f"Wrote {doc_path}")
    print(f"Wrote {gt_path}  ({len(ground_truth)} ground-truth entities)")


if __name__ == "__main__":
    build()
