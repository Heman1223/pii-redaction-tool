"""End-to-end tests: replacement consistency, tables, and DOCX integrity.

These cover the requirements that cannot be checked by looking at one string:
the same entity must map to the same fake value everywhere, two people sharing
a name must not be merged, table cells must be processed, and the written file
must still be a valid, openable DOCX.
"""

from pathlib import Path

import pytest
from docx import Document

from src.document_utils import iter_text_units, load_document
from src.evaluator import CategoryMetrics, evaluate
from src.redactor import redact_document
from src.replacements import ReplacementMap


# ---------------------------------------------------------------------------
# Consistent replacement
# ---------------------------------------------------------------------------


def test_same_entity_gets_same_fake_value():
    """The central requirement: repeated PII must redact identically."""
    mapping = ReplacementMap(seed=1)
    first = mapping.get("PERSON", "Rashi Patil")
    second = mapping.get("PERSON", "Rashi Patil")
    assert first == second


def test_case_and_spacing_variants_share_one_value():
    """Cover pages shout names in capitals; the body uses title case."""
    mapping = ReplacementMap(seed=1)
    assert mapping.get("PERSON", "RAJESH  KUSHAL HEGDE") == mapping.get(
        "PERSON", "Rajesh Kushal Hegde"
    )


def test_phone_formats_share_one_fake_identity():
    """The same number written two ways must redact to one fake number.

    A document prints a number as "+91 98765 43210" in a contact block and
    "+91-9876543210" in a footer. Separators are cosmetic, so both mentions are
    one phone number and must receive one fake identity - otherwise the
    redacted document invents a second contact that never existed.

    Regression: identity was previously keyed on whitespace/case normalisation,
    which leaves separators intact, so the two spellings produced two fakes.
    """
    mapping = ReplacementMap(seed=42)
    spaced = mapping.get("PHONE", "+91 98765 43210")
    dashed = mapping.get("PHONE", "+91-9876543210")
    assert spaced == dashed


def test_phone_identity_ignores_all_separator_styles():
    """Every separator style of one number collapses to a single identity."""
    mapping = ReplacementMap(seed=42)
    variants = [
        "+91 98765 43210",
        "+91-9876543210",
        "+91 9876543210",
        "+91-98765-43210",
        "+91  98765  43210",
    ]
    assert len({mapping.get("PHONE", v) for v in variants}) == 1


def test_distinct_phone_numbers_still_differ():
    """Canonicalisation must not collapse two genuinely different numbers."""
    mapping = ReplacementMap(seed=42)
    assert mapping.get("PHONE", "+91 98765 43210") != mapping.get(
        "PHONE", "+91 22 4009 4400"
    )


def test_mapping_report_shows_the_original_spelling():
    """The canonical key is digits; the report must stay human-readable."""
    mapping = ReplacementMap(seed=42)
    mapping.get("PHONE", "+91 98765 43210")
    keys = list(mapping.as_dict())
    assert any("+91 98765 43210" in k for k in keys), keys


def test_different_entities_get_different_values():
    mapping = ReplacementMap(seed=1)
    assert mapping.get("PERSON", "Rashi Patil") != mapping.get("PERSON", "Rohan Dey")


def test_same_name_different_role_gets_different_identity():
    """The same-name edge case from the brief.

    Two different people both called "Rahul Sharma" are told apart by the role
    stated beside each one, so they receive separate fake identities.
    """
    mapping = ReplacementMap(seed=1)
    director = mapping.get("PERSON", "Rahul Sharma", context="director")
    cfo = mapping.get("PERSON", "Rahul Sharma", context="chief financial officer")
    assert director != cfo


def test_same_name_same_role_is_merged_known_limitation():
    """Documented limitation: identical name AND role cannot be separated.

    With no distinguishing context there is no signal to split on, so the two
    people are merged. This test pins the behaviour so it stays a known,
    deliberate trade rather than an accident.
    """
    mapping = ReplacementMap(seed=1)
    a = mapping.get("PERSON", "Rahul Sharma", context="director")
    b = mapping.get("PERSON", "Rahul Sharma", context="director")
    assert a == b


def test_replacements_are_type_compatible():
    mapping = ReplacementMap(seed=7)
    assert "@" in mapping.get("EMAIL", "someone@example.org")
    assert mapping.get("PHONE", "+91 9876543210").startswith("+91")
    assert mapping.get("IP_ADDRESS", "10.0.0.1").startswith("192.0.2.")
    assert mapping.get("SSN", "123-45-6789").count("-") == 2


def test_replacements_are_deterministic_across_runs():
    """A fixed seed keeps the evaluation report reproducible."""
    a = ReplacementMap(seed=42).get("PERSON", "Rashi Patil")
    b = ReplacementMap(seed=42).get("PERSON", "Rashi Patil")
    assert a == b


# ---------------------------------------------------------------------------
# Document processing
# ---------------------------------------------------------------------------


@pytest.fixture
def docx_with_table(tmp_path) -> Path:
    """A document whose PII lives only inside a table."""
    doc = Document()
    doc.add_paragraph("Directory of contacts follows.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Email"
    table.cell(1, 0).text = "Priya Nair"
    table.cell(1, 1).text = "priya.nair@example.org"

    path = tmp_path / "table.docx"
    doc.save(path)
    return path


def test_table_cells_are_processed(docx_with_table, tmp_path):
    """PII inside tables must be redacted, not just document.paragraphs."""
    out = tmp_path / "table_redacted.docx"
    result = redact_document(str(docx_with_table), str(out))

    assert "EMAIL" in result.counts_by_type

    text = "\n".join(u.text for u in iter_text_units(load_document(str(out))))
    assert "priya.nair@example.org" not in text


def test_output_is_a_valid_openable_docx(docx_with_table, tmp_path):
    out = tmp_path / "valid.docx"
    redact_document(str(docx_with_table), str(out))
    assert out.exists() and out.stat().st_size > 0
    Document(str(out))  # raises if the file is corrupt


def test_original_document_is_not_modified(docx_with_table, tmp_path):
    before = docx_with_table.read_bytes()
    redact_document(str(docx_with_table), str(tmp_path / "copy.docx"))
    assert docx_with_table.read_bytes() == before


def test_consistency_across_the_whole_document(tmp_path):
    """One person mentioned three times must redact to one fake name."""
    doc = Document()
    doc.add_paragraph("Rashi Patil joined the team.")
    doc.add_paragraph("Later, Rashi Patil was promoted.")
    doc.add_paragraph("Finally Rashi Patil resigned.")
    src = tmp_path / "repeat.docx"
    doc.save(src)

    out = tmp_path / "repeat_redacted.docx"
    result = redact_document(str(src), str(out))

    text = "\n".join(u.text for u in iter_text_units(load_document(str(out))))
    assert "Rashi Patil" not in text

    # Exactly one person was seen, so exactly one fake identity was created...
    person_fakes = {
        fake for key, fake in result.replacement_map.items()
        if key.startswith("PERSON:")
    }
    assert len(person_fakes) == 1

    # ...and it must appear in all three sentences.
    assert text.count(person_fakes.pop()) == 3


def test_images_are_neutralised(tmp_path):
    """Images cannot be read without OCR, so they are blanked wholesale.

    The reference prospectus hides a PAN card - name, father's name, date of
    birth and PAN number - inside an image where no text rule can reach it.
    """
    from docx.shared import Inches
    import base64, io

    png = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNk"
        "YPhfDwAChwGA60e6kgAAAABJRU5ErkJggg=="
    )
    doc = Document()
    doc.add_paragraph("See the attached identity document.")
    doc.add_picture(io.BytesIO(png), width=Inches(1))
    src = tmp_path / "with_image.docx"
    doc.save(src)

    result = redact_document(str(src), str(tmp_path / "img_redacted.docx"))
    assert result.images_found >= 1
    assert result.images_blanked == result.images_found


# ---------------------------------------------------------------------------
# Evaluator
# ---------------------------------------------------------------------------


def test_metrics_arithmetic():
    m = CategoryMetrics(tp=8, fp=2, fn=2)
    assert m.precision == pytest.approx(0.8)
    assert m.recall == pytest.approx(0.8)
    assert m.f1 == pytest.approx(0.8)
    assert m.accuracy == pytest.approx(8 / 12)   # TP / (TP + FP + FN)


def test_metrics_are_zero_when_nothing_predicted():
    m = CategoryMetrics(tp=0, fp=0, fn=5)
    assert m.precision == 0.0 and m.recall == 0.0 and m.f1 == 0.0


def test_phone_formatting_differences_still_match():
    report = evaluate(
        [("+91 20 4505 3237", "PHONE")], [("+91 20 45053237", "PHONE")]
    )
    assert report.overall.tp == 1 and report.overall.fp == 0


def test_dash_style_differences_still_match():
    """En-dash vs hyphen must not be scored as an error."""
    report = evaluate(
        [("Pune – 411 044, Maharashtra, India", "ADDRESS")],
        [("Pune - 411 044, Maharashtra, India", "ADDRESS")],
    )
    assert report.overall.tp == 1 and report.overall.fp == 0


def test_address_fragments_count_as_one_detection():
    """A multi-paragraph address split into fragments is one true positive."""
    report = evaluate(
        [("12 Example Road", "ADDRESS"), ("Pune - 411 001", "ADDRESS")],
        [("12 Example Road, Pune - 411 001", "ADDRESS")],
    )
    assert report.overall.tp == 1
    assert report.overall.fp == 0


def test_wrong_type_is_not_a_match():
    report = evaluate([("Acme Ltd", "PERSON")], [("Acme Ltd", "COMPANY")])
    assert report.overall.tp == 0
    assert report.overall.fp == 1 and report.overall.fn == 1


# ---------------------------------------------------------------------------
# Web app: retention and UI/backend consistency
# ---------------------------------------------------------------------------


def test_expired_job_directories_are_purged(tmp_path, monkeypatch):
    """Uploads must not linger, since they hold the PII being redacted.

    This is what makes the retention note shown in the UI a true statement
    rather than a marketing claim.
    """
    import os, time as _time
    from src import app as webapp

    monkeypatch.setattr(webapp, "JOBS_DIR", tmp_path)
    monkeypatch.setattr(webapp, "JOB_RETENTION_SECONDS", 3600)

    old = tmp_path / "oldjob"
    old.mkdir()
    (old / "secret.docx").write_bytes(b"PK")
    # Backdate it well past the retention window.
    stale = _time.time() - 7200
    os.utime(old, (stale, stale))

    fresh = tmp_path / "freshjob"
    fresh.mkdir()
    (fresh / "secret.docx").write_bytes(b"PK")

    webapp._purge_expired_jobs()

    assert not old.exists(), "expired job files must be deleted"
    assert fresh.exists(), "a recent job must be kept so its download still works"


def test_ui_does_not_hardcode_the_upload_limit():
    """The landing page must render the backend's real limit.

    A hardcoded number silently becomes a lie the moment the backend changes,
    promising a size the server will reject.
    """
    from pathlib import Path

    template = Path("templates/index.html").read_text(encoding="utf-8")
    assert "{{ max_upload_mb }}" in template
    assert "{{ accept_attr }}" in template
    assert "50 MB" not in template and "50&nbsp;MB" not in template


def test_accepted_formats_match_the_backend():
    """Every format offered in the file picker must actually be processed."""
    from src.app import ACCEPT_ATTR, DOCX_SUFFIXES, TEXT_SUFFIXES

    offered = set(ACCEPT_ATTR.split(","))
    assert offered == DOCX_SUFFIXES | TEXT_SUFFIXES
