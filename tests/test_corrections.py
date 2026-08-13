"""Regression tests for the correction pass.

Each test here pins a defect that was found by running the tool on adversarial
input, so the behaviour cannot silently regress. Grouped by the issue that
produced them.
"""

import tempfile
from pathlib import Path

import pytest
from docx import Document

from src.detector import detect_address_entities, detect_entities
from src.document_utils import document_from_text, iter_text_units, load_document
from src.redactor import (
    _contexts_conflict,
    _split_context,
    redact_document,
    resolve_person_identities,
)
from src.detector import Entity


# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------


def addresses(text):
    return [e.text for e in detect_address_entities(text)]


def redact_text(tmp_path, text, name="doc"):
    """Redact a block of plain text and return (output lines, result)."""
    source = tmp_path / f"{name}.docx"
    document_from_text(text, str(source))
    output = tmp_path / f"{name}_redacted.docx"
    result = redact_document(str(source), str(output))
    lines = [u.text for u in iter_text_units(load_document(str(output)))]
    return lines, result


def fakes_of(result, prefix):
    return {v for k, v in result.replacement_map.items() if k.startswith(prefix)}


# ===========================================================================
# 1. Address span boundaries
# ===========================================================================

ADDRESS = "106 Example Street, Sector 7, Exampleton - 411 006, Example State, India"


@pytest.mark.parametrize("trailing", [
    "Order ID: ORD-458293",
    "Company Code KSH-2026-4812",
    "Invoice: INV-2026-1024",
    "Ticket ID: TKT-2026-08421",
    "Request ID: REQ-987654321",
    "and the meeting was held there.",
])
def test_address_span_stops_at_the_address(trailing):
    """The span must cover the address and nothing after it.

    The old rule ended with "(?:\\s*,?\\s*[A-Z][A-Za-z.]+){0,3}", which grabbed
    any three capitalised words after the postcode - so "... India Ticket ID:
    TKT-2026-08421" was swallowed whole and destroyed by the replacement.
    """
    found = addresses(f"{ADDRESS} {trailing}")
    assert found, "the address itself must still be detected"
    assert found[0] == ADDRESS


def test_address_followed_by_id_does_not_absorb_it():
    text = f"{ADDRESS}-458293, Company Code KSH-2026-4812"
    found = addresses(text)
    assert found == [ADDRESS]
    assert "458293" not in found[0]
    assert "KSH-2026-4812" not in found[0]


def test_trailing_full_stop_is_not_part_of_the_address():
    found = addresses(f"{ADDRESS}. The office is open on weekdays.")
    assert found == [ADDRESS]


def test_order_id_alone_is_not_an_address():
    """"ORD-458293" ends in six digits, exactly like an Indian PIN code."""
    assert addresses("Order ID: ORD-458293") == []
    assert addresses("Reference REQ-987654321 was raised") == []


def test_multiline_address_is_still_detected(tmp_path):
    """The boundary fix must not break addresses split across paragraphs."""
    lines, result = redact_text(
        tmp_path,
        "KSH International Limited\n"
        "11/3, 11/4 and 11/5, Village Birdewadi Chakan, Taluka-Khed\n"
        "Pune - 410 501\n"
        "Maharashtra, India",
        "multiline",
    )
    assert "ADDRESS" in result.counts_by_type
    joined = "\n".join(lines)
    assert "Village Birdewadi" not in joined
    assert "410 501" not in joined


# ===========================================================================
# 2. Non-PII adjacent to addresses must survive
# ===========================================================================

ADVERSARIAL = f"""Registered Office:
{ADDRESS}
Company Code: KSH-2026-4812
Order ID: ORD-458293
Ticket ID: TKT-2026-08421
Invoice: INV-2026-1024
Request ID: REQ-987654321
Status code: 200
Date: 13 August 2026
Amount: 45,000"""

PRESERVED = [
    "KSH-2026-4812", "ORD-458293", "TKT-2026-08421", "INV-2026-1024",
    "REQ-987654321", "Status code: 200", "13 August 2026", "45,000",
]


@pytest.mark.parametrize("token", PRESERVED)
def test_identifiers_next_to_an_address_are_preserved(tmp_path, token):
    lines, _ = redact_text(tmp_path, ADVERSARIAL, "adversarial")
    assert token in "\n".join(lines)


def test_adversarial_block_redacts_only_the_address(tmp_path):
    lines, result = redact_text(tmp_path, ADVERSARIAL, "adv_only")
    joined = "\n".join(lines)
    assert "Example Street" not in joined or "106 Example Street" not in joined
    assert result.counts_by_type.get("ADDRESS", 0) >= 1
    # Nothing else in this block is PII.
    assert set(result.counts_by_type) <= {"ADDRESS"}


# ===========================================================================
# 3. Same-person consistency across every type
# ===========================================================================


def test_repeated_entities_all_receive_identical_replacements(tmp_path):
    """Every repeated entity must map to exactly one fake value."""
    block = (
        "Rahul Sharma, Director, Acme Technologies Ltd.\n"
        "Email rahul.sharma@acme.com or call +91 9876543210.\n"
        "Office: 42 Baker Street, Kothrud, Pune - 411 038, Maharashtra, India\n"
    )
    lines, result = redact_text(tmp_path, block * 3, "repeat")

    for prefix in ("PERSON:", "EMAIL:", "PHONE:", "COMPANY:"):
        assert len(fakes_of(result, prefix)) == 1, f"{prefix} was not consistent"

    joined = "\n".join(lines)
    for original in ("Rahul Sharma", "rahul.sharma@acme.com",
                     "+91 9876543210", "Acme Technologies", "Baker Street"):
        assert original not in joined

    # The one fake person must appear once per original mention.
    person = fakes_of(result, "PERSON:").pop()
    assert joined.count(person) == 3


def test_case_variants_share_one_identity(tmp_path):
    lines, result = redact_text(
        tmp_path,
        "RAHUL SHARMA signed the register.\nRahul Sharma countersigned it.",
        "casevariants",
    )
    assert len(fakes_of(result, "PERSON:")) == 1


# ===========================================================================
# 4. Same name, different person
# ===========================================================================


def test_case_a_same_name_same_context_is_one_person(tmp_path):
    _, result = redact_text(
        tmp_path,
        "Rahul Sharma, Director, Acme Technologies Ltd. approved it.\n"
        "Later Rahul Sharma, Director, Acme Technologies Ltd. signed it.",
        "case_a",
    )
    assert len(fakes_of(result, "PERSON:")) == 1


def test_case_b_same_name_different_role_and_org_is_two_people(tmp_path):
    _, result = redact_text(
        tmp_path,
        "Rahul Sharma, Director, Acme Technologies Ltd. approved it.\n"
        "Rahul Sharma, Senior Manager, Beta Financial Services Limited, objected.",
        "case_b",
    )
    assert len(fakes_of(result, "PERSON:")) == 2


def test_case_c_insufficient_context_is_deterministic(tmp_path):
    """Documented limitation: with no evidence, mentions are merged.

    Merging is the safer default - splitting one real person into several fake
    identities makes the redacted document incoherent, which is a worse failure
    than merging two strangers who happen to share a name.
    """
    _, first = redact_text(tmp_path, "Rahul Sharma attended.\nRahul Sharma left.", "case_c1")
    _, second = redact_text(tmp_path, "Rahul Sharma attended.\nRahul Sharma left.", "case_c2")

    assert len(fakes_of(first, "PERSON:")) == 1
    assert fakes_of(first, "PERSON:") == fakes_of(second, "PERSON:")


def test_partial_job_title_is_not_a_different_person():
    """"CS and Compliance Officer" abbreviates "Company Secretary and ...".

    Treating an abbreviation as a conflicting role split one real person into
    two identities on the reference document.
    """
    full = _split_context("company secretary and compliance officer||")
    short = _split_context("compliance officer||")
    assert not _contexts_conflict(full, short)


def test_genuinely_different_roles_do_conflict():
    assert _contexts_conflict(_split_context("director||"),
                              _split_context("senior manager||"))


def test_same_role_different_org_conflicts():
    assert _contexts_conflict(_split_context("director||acme technologies ltd."),
                              _split_context("director||beta financial services limited"))


def test_organisation_alone_does_not_split_a_person():
    """Organisation is proximity evidence and is deliberately weak.

    A paragraph merely mentioning a different company is not a statement that
    this is a different person; treating it as one split a single promoter into
    four identities on the reference document.
    """
    assert not _contexts_conflict(_split_context("||acme technologies ltd."),
                                  _split_context("||beta financial services limited"))


def test_resolver_merges_when_only_one_context_cluster():
    entities = [
        Entity("Rahul Sharma", "PERSON", 0, 12, context="director||"),
        Entity("Rahul Sharma", "PERSON", 0, 12, context=None),
        Entity("Rahul Sharma", "PERSON", 0, 12, context="director||acme ltd."),
    ]
    resolved = resolve_person_identities(entities)
    assert {e.context for e in resolved} == {None}


# ===========================================================================
# 5. Overlapping entities
# ===========================================================================


def test_email_does_not_also_produce_a_person(tmp_path):
    text = "Write to john.smith@example.com about the order."
    entities = detect_entities(text)
    emails = [e for e in entities if e.type == "EMAIL"]
    assert len(emails) == 1
    for other in entities:
        if other is not emails[0]:
            assert not other.overlaps(emails[0])


def test_numbers_inside_an_address_do_not_become_other_pii():
    entities = detect_entities(ADDRESS, use_ner=False)
    assert {e.type for e in entities} <= {"ADDRESS"}


def test_address_and_following_id_do_not_overlap():
    entities = detect_entities(f"{ADDRESS} Order ID: ORD-458293", use_ner=False)
    for a in entities:
        for b in entities:
            if a is not b:
                assert not a.overlaps(b)


# ===========================================================================
# 6. Non-PII regression
# ===========================================================================


@pytest.mark.parametrize("line", [
    "Ticket ID: TKT-2026-08421",
    "Order ID: ORD-458293",
    "Invoice: INV-2026-1024",
    "Company Code: KSH-2026-4812",
    "Request ID: REQ-987654321",
    "Status code: 200",
    "Date: 13 August 2026",
    "Amount: 45,000",
    "Corporate identity number: U28129PN1979PLC141032",
    "SEBI Registration Number: INM000013004",
    "Total revenue was 19,282.93 million for Fiscal 2025",
    "Shares allotted: 26,704,570 at face value 5 each",
])
def test_non_pii_lines_are_left_alone(line):
    assert detect_entities(line, use_ner=False) == []


# ===========================================================================
# 7. Tables get the same treatment as paragraphs
# ===========================================================================


def _table_doc(path, rows):
    doc = Document()
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            table.cell(r, c).text = value
    doc.save(path)


def test_address_boundaries_hold_inside_table_cells(tmp_path):
    """The boundary fix must apply to tables, not just paragraphs."""
    source = tmp_path / "table_addr.docx"
    _table_doc(source, [
        ("Office", "Reference"),
        (f"{ADDRESS} Order ID: ORD-458293", "Company Code: KSH-2026-4812"),
    ])
    output = tmp_path / "table_addr_redacted.docx"
    redact_document(str(source), str(output))

    text = "\n".join(u.text for u in iter_text_units(load_document(str(output))))
    assert "ORD-458293" in text
    assert "KSH-2026-4812" in text
    assert "106 Example Street" not in text


def test_consistency_between_a_table_cell_and_a_paragraph(tmp_path):
    """The same person named in prose and in a table gets one identity."""
    doc = Document()
    doc.add_paragraph("Rahul Sharma, Director, Acme Technologies Ltd. approved it.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Role"
    table.cell(1, 0).text = "Rahul Sharma"
    table.cell(1, 1).text = "Director"

    source = tmp_path / "mixed.docx"
    doc.save(source)
    output = tmp_path / "mixed_redacted.docx"
    result = redact_document(str(source), str(output))

    assert len(fakes_of(result, "PERSON:")) == 1
    text = "\n".join(u.text for u in iter_text_units(load_document(str(output))))
    assert "Rahul Sharma" not in text


# ===========================================================================
# 8. Image handling is unchanged
# ===========================================================================


def test_images_are_still_neutralised_without_ocr(tmp_path):
    """Images are blanked because they cannot be inspected - not "detected"."""
    import base64
    import io

    from docx.shared import Inches

    png = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNk"
        "YPhfDwAChwGA60e6kgAAAABJRU5ErkJggg=="
    )
    doc = Document()
    doc.add_paragraph("Identity document attached.")
    doc.add_picture(io.BytesIO(png), width=Inches(1))
    source = tmp_path / "img.docx"
    doc.save(source)

    result = redact_document(str(source), str(tmp_path / "img_out.docx"))
    assert result.images_found >= 1
    assert result.images_blanked == result.images_found
